"""
Word (DOCX) 渲染模块
=====================

将优化后的简历内容渲染为可编辑的 Word (.docx) 文档。

功能:
- ATS 兼容格式（纯文本、无页眉/页脚/文本框）
- 标准简历排版（标题层级、项目符号、粗体关键词）
- 支持中文（使用系统默认字体）
- 可在 Word/WPS/Google Docs 中直接编辑

依赖: python-docx (已安装)
"""

from __future__ import annotations
import os
import logging
import tempfile
from typing import Optional

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from .interfaces import (
    OptimizationResult,
    ResumeSection,
    DOCXOutput,
)

logger = logging.getLogger(__name__)


def _set_east_asia_font(run, font_name: str = '黑体') -> None:
    """安全设置中文字体，处理 rPr 可能为 None 的情况。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def _format_section_content_for_docx(doc: Document, content: str) -> None:
    """
    将段落内容添加为格式化文本。
    处理项目符号列表和关键词高亮。
    """
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')  # 空行
            continue
        
        # 项目符号
        if stripped.startswith('•') or stripped.startswith('- ') or stripped.startswith('* '):
            item_text = stripped.lstrip('•-').strip()
            p = doc.add_paragraph(item_text, style='List Bullet')
            # 加粗关键词
            _highlight_paragraph_keywords(p)
        else:
            p = doc.add_paragraph(stripped)
            _highlight_paragraph_keywords(p)


def _highlight_paragraph_keywords(paragraph) -> None:
    """
    对段落中的技术关键词加粗。
    直接在 run 级别设置 bold。
    """
    tech_keywords = [
        'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'Go', 'Rust',
        'React', 'Vue', 'Angular', 'Spring', 'Django', 'Flask', 'FastAPI',
        'TensorFlow', 'PyTorch', 'Docker', 'Kubernetes', 'AWS', 'Azure',
        'MySQL', 'PostgreSQL', 'Redis', 'MongoDB', 'Git', 'Linux',
        'AI', 'ML', 'NLP', 'LLM',
    ]
    for run in paragraph.runs:
        text = run.text
        for kw in tech_keywords:
            if kw in text:
                run.bold = True
                break


def _add_heading_styled(doc: Document, text: str, level: int, size: int) -> None:
    """添加带中文字体的标题。"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x2C)
        _set_east_asia_font(run, '黑体')
    return heading


def _add_run_with_font(paragraph, text: str, bold: bool = False, size: int = 10, color: Optional[tuple] = None) -> None:
    """添加带中文字体的 run。"""
    run = paragraph.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    _set_east_asia_font(run, '黑体')


def render_docx(
    optimization_result: OptimizationResult,
    output_path: Optional[str] = None,
    job_title: Optional[str] = None,
    company_name: Optional[str] = None,
) -> DOCXOutput:
    """
    将优化结果渲染为可编辑的 Word (.docx) 文档。

    Args:
        optimization_result: 优化后的简历段落
        output_path: 输出路径（默认临时文件）
        job_title: 目标职位名称（用于文档标题）
        company_name: 目标公司名称

    Returns:
        DOCXOutput 包含文件信息

    Raises:
        DOCXGenerationError: 文档生成失败
    """
    logger.info("Rendering DOCX, job_title=%s, company=%s", job_title, company_name)

    if not output_path:
        fd, output_path = tempfile.mkstemp(suffix='.docx', prefix='resumefit_')
        os.close(fd)

    doc = Document()

    # ============================================================
    # 配置页面边距
    # ============================================================
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ============================================================
    # 设置默认字体（支持中文）
    # ============================================================
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _set_east_asia_font(font, '黑体')

    # ============================================================
    # 简历标题
    # ============================================================
    title_text = '个人简历'
    if job_title:
        title_text += f' — 应聘 {job_title}'
    if company_name:
        title_text += f' @ {company_name}'
    
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        _set_east_asia_font(run, '黑体')

    # 分隔线
    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = divider.add_run('━' * 40)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(8)

    # ============================================================
    # 简历段落
    # ============================================================
    for section in optimization_result.sections:
        # 段落标题
        _add_heading_styled(doc, section.title, level=2, size=14)
        
        # 段落内容
        _format_section_content_for_docx(doc, section.content)

    # ============================================================
    # 变更说明（附加部分）
    # ============================================================
    if optimization_result.changes:
        doc.add_page_break()
        _add_heading_styled(doc, '📝 优化变更说明', level=1, size=16)
        
        changes_note = doc.add_paragraph()
        _add_run_with_font(changes_note, '以下为本次简历优化的变更明细，供参考和手动调整。', size=10, color=(0x66, 0x66, 0x66))
        
        for change in optimization_result.changes:
            _add_heading_styled(doc, f'「{change.section_title}」', level=3, size=12)
            
            p_type = doc.add_paragraph()
            _add_run_with_font(p_type, f'变更类型: {change.change_type}', bold=True, size=10)
            
            p_reason = doc.add_paragraph()
            _add_run_with_font(p_reason, f'原因: {change.reason}', size=10)
            
            p_before = doc.add_paragraph()
            _add_run_with_font(p_before, '优化前: ', bold=True, size=10)
            _add_run_with_font(p_before, change.original_text, size=10)
            
            p_after = doc.add_paragraph()
            _add_run_with_font(p_after, '优化后: ', bold=True, size=10)
            _add_run_with_font(p_after, change.optimized_text, size=10)
            
            doc.add_paragraph('')  # 分隔空行

    # ============================================================
    # 保存
    # ============================================================
    doc.save(output_path)
    file_size = os.path.getsize(output_path)
    
    logger.info("DOCX saved to %s, size=%d bytes", output_path, file_size)
    
    return DOCXOutput(
        file_path=output_path,
        file_size_bytes=file_size,
        ats_compatible=True,
        editable=True,
    )
